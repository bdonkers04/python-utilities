from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches
import os

"""
Excel-to-Word Data & Image Migrator
This script automates the extraction of tabular data and embedded imagery 
from an Excel (.xlsx) spreadsheet into a structured Word (.docx) document.
"""

# 1. DATA EXTRACTION PHASE
# Initialize the openpyxl engine to read the source spreadsheet
wb = load_workbook("original.xlsx")
ws = wb.active  # Target the primary/active worksheet

# 2. DOCUMENT ARCHITECTURE PHASE
# Initialize a new python-docx object and set the main title
doc = Document()
doc.add_heading("Products", level=1)

# Dynamically determine the grid size based on the Excel content
rows = ws.max_row
cols = ws.max_column

# Pre-allocate a table in the Word document with matching dimensions
table = doc.add_table(rows=rows, cols=cols)

# 3. CONTENT MIGRATION (TEXT)
# Iterate through Excel rows and columns to map values to the Word table
# iter_rows(values_only=True) is used for performance to skip cell styling
for i, row in enumerate(ws.iter_rows(values_only=True)):
    for j, val in enumerate(row):
        # Assign text to the corresponding cell, handling empty (None) values
        table.cell(i, j).text = str(val) if val is not None else ""

# 4. MULTIMEDIA HANDLING (IMAGES)
# openpyxl stores images in a non-standard '_images' attribute
if hasattr(ws, "_images"): 
    for img in ws._images:
        # Define a temporary bridge file for the image binary
        img_path = "temp_img.png"
        
        # Calculate a rough positional reference from the Excel anchor
        img.ref = img.anchor._from.col + 1 
        
        # Extract the binary image data and save it to the local disk
        img.image.save(img_path)

        # Append the image to the end of the Word document
        # Set a fixed width of 1.5 inches to maintain visual consistency
        doc.add_picture(img_path, width=Inches(1.5))

        # Cleanup: Remove the temporary file to keep the directory clean
        os.remove(img_path)

# 5. FINALIZATION
# Write the generated document to the file system
doc.save("products.docx")
